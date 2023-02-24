import datetime
import math
import docx

def create_docx_file(data, output_filename):
    document = docx.Document()
    for i, item in enumerate(data):
        start = datetime.timedelta(seconds=item['start'])
        end = datetime.timedelta(seconds=item['end'])
        start_srt = start_to_srt(start)
        end_srt = start_to_srt(end)
        paragraph = document.add_paragraph()
        paragraph.add_run(str(i+1) + '\n').bold = True
        paragraph.add_run(start_srt + ' --> ' + end_srt + '\n').italic = True
        paragraph.add_run(item['text'] + '\n')
        paragraph.add_run('\n')
    document.save(output_filename)

def create_srt_file(data, output_filename):
    with open(output_filename, 'w') as f:
        for i, item in enumerate(data):
            start = datetime.timedelta(seconds=item['start'])
            end = datetime.timedelta(seconds=item['end'])
            start_srt = start_to_srt(start)
            end_srt = start_to_srt(end)
            f.write(str(i+1) + '\n')
            f.write(start_srt + ' --> ' + end_srt + '\n')
            f.write(item['text'] + '\n')
            f.write('\n')

def start_to_srt(start):
    hours, remainder = divmod(start.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    milliseconds = math.floor(start.microseconds / 1000)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"

data = [{'id': 0, 'seek': 0, 'start': 0.0, 'end': 9.48, 'text': ' Okay, so previously we have covered in biology the blood system, how we had a bunch of questions', 'tokens': [1033, 11, 370, 8046, 321, 362, 5343, 294, 14956, 264, 3390, 1185, 11, 577, 321, 632, 257, 3840, 295, 1651], 'temperature': 0.0, 'avg_logprob': -0.15691421071036918, 'compression_ratio': 1.8646616541353382, 'no_speech_prob': 0.1497831642627716}, {'id': 1, 'seek': 0, 'start': 9.48, 'end': 16.36, 'text': ' that we could create about the blood system, and we had questions that were, you know, asking', 'tokens': [300, 321, 727, 1884, 466, 264, 3390, 1185, 11, 293, 321, 632, 1651, 300, 645, 11, 291, 458, 11, 3365], 'temperature': 0.0, 'avg_logprob': -0.15691421071036918, 'compression_ratio': 1.8646616541353382, 'no_speech_prob': 0.1497831642627716}, {'id': 2, 'seek': 0, 'start': 16.36, 'end': 19.84, 'text': ' every single point, what is it, and why is it important.', 'tokens': [633, 2167, 935, 11, 437, 307, 309, 11, 293, 983, 307, 309, 1021, 13], 'temperature': 0.0, 'avg_logprob': -0.15691421071036918, 'compression_ratio': 1.8646616541353382, 'no_speech_prob': 0.1497831642627716}, {'id': 3, 'seek': 0, 'start': 19.84, 'end': 23.080000000000002, 'text': " And remember, we're asking these for our headings as well, why are vessels important, why is", 'tokens': [400, 1604, 11, 321, 434, 3365, 613, 337, 527, 1378, 1109, 382, 731, 11, 983, 366, 20117, 1021, 11, 983, 307], 'temperature': 0.0, 'avg_logprob': -0.15691421071036918, 'compression_ratio': 1.8646616541353382, 'no_speech_prob': 0.1497831642627716}]

create_srt_file(data, "test.srt")
create_docx_file(data, "test.docx")




